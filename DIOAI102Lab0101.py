import os
import requests
from docx import Document
from typing import List, Optional
from dotenv import load_dotenv

load_dotenv("DIOLab01.env")

# Load environment variables
AZURE_KEY = os.getenv("AZURE_TRANSLATOR_KEY")
AZURE_ENDPOINT = os.getenv("AZURE_ENDPOINT", "https://api.cognitive.microsofttranslator.com")
AZURE_LOCATION = os.getenv("AZURE_LOCATION", "eastus2")

if not AZURE_KEY:
    raise ValueError("AZURE_TRANSLATOR_KEY environment variable not set")

class DocumentTranslator:
    """Handles professional document translation using Azure Cognitive Services."""
    
    def __init__(self, target_lang: str = "pt-br"):
        self.target_lang = target_lang
        self.base_url = f"{AZURE_ENDPOINT}/translate"
        self.headers = {
            'Ocp-Apim-Subscription-Key': AZURE_KEY,
            'Ocp-Apim-Subscription-Region': AZURE_LOCATION,
            'Content-type': 'application/json',
        }

    def translate_batch(self, texts: List[str]) -> List[str]:
        """Translates a list of strings in a single API call for efficiency."""
        # Azure allows up to 100 elements per batch
        params = {'api-version': '3.0', 'from': 'en', 'to': self.target_lang}
        body = [{'text': text} for text in texts if text.strip()]
        
        if not body:
            return []

        try:
            response = requests.post(self.base_url, params=params, headers=self.headers, json=body)
            response.raise_for_status()
            data = response.json()
            return [item["translations"][0]["text"] for item in data]
        except Exception as e:
            print(f"Translation Error: {e}")
            return texts

    def process_document(self, input_path: str) -> str:
        """Reads, translates, and saves a new document."""
        if not os.path.exists(input_path):
            raise FileNotFoundError(f"File not found: {input_path}")

        doc = Document(input_path)
        
        paragraphs_to_translate = [p.text for p in doc.paragraphs if p.text.strip()]
        
        translated_texts = self.translate_batch(paragraphs_to_translate)

        new_doc = Document()
        for text in translated_texts:
            new_doc.add_paragraph(text)

        output_path = input_path.replace(".docx", f"_{self.target_lang}.docx")
        new_doc.save(output_path)
        return output_path

if __name__ == "__main__":
    translator = DocumentTranslator(target_lang="pt-br")
    try:
        path = "Niccolo Machiavelli by Matthew Wills.docx"
        result = translator.process_document(path)
        print(f"Success! Document saved at: {result}")
    except Exception as e:
        print(f"Critical Failure: {e}")